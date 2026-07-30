import os
import fitz  # PyMuPDF
import cv2
import numpy as np
from PIL import Image
import zipfile
from datetime import datetime
from docx import Document
from docx.shared import Inches
import io
import re

class PDFProcessor:
    def __init__(self, upload_folder='uploads', processed_folder='processed'):
        self.upload_folder = upload_folder
        self.processed_folder = processed_folder
    
    def remove_header_and_seal(self, session_id, files):
        """去红头及公章"""
        try:
            processed_files = []
            processed_dir = os.path.join(self.processed_folder, session_id)
            os.makedirs(processed_dir, exist_ok=True)
            
            for file_info in files:
                if file_info['type'] in ['文本型PDF', '扫描型PDF']:
                    input_path = os.path.join(self.upload_folder, session_id, file_info['filename'])
                    output_filename = f"去红头-{file_info['filename']}"
                    output_path = os.path.join(processed_dir, output_filename)
                    
                    if file_info['type'] == '文本型PDF':
                        self._remove_header_seal_text_pdf(input_path, output_path)
                    else:
                        self._remove_header_seal_scan_pdf(input_path, output_path)
                    
                    processed_files.append(output_filename)
            
            if len(processed_files) == 1:
                return {'success': True, 'download_file': processed_files[0]}
            elif len(processed_files) > 1:
                zip_filename = '去红头.zip'
                zip_path = os.path.join(processed_dir, zip_filename)
                self._create_zip(processed_dir, processed_files, zip_path)
                return {'success': True, 'download_file': zip_filename}
            else:
                return {'error': '没有可处理的PDF文件'}
                
        except Exception as e:
            return {'error': f'处理失败: {str(e)}'}
    
    def remove_seal_only(self, session_id, files):
        """仅去公章"""
        try:
            processed_files = []
            processed_dir = os.path.join(self.processed_folder, session_id)
            os.makedirs(processed_dir, exist_ok=True)
            
            for file_info in files:
                if file_info['type'] in ['文本型PDF', '扫描型PDF']:
                    input_path = os.path.join(self.upload_folder, session_id, file_info['filename'])
                    output_filename = f"去公章-{file_info['filename']}"
                    output_path = os.path.join(processed_dir, output_filename)
                    
                    if file_info['type'] == '文本型PDF':
                        self._remove_seal_text_pdf(input_path, output_path)
                    else:
                        self._remove_seal_scan_pdf(input_path, output_path)
                    
                    processed_files.append(output_filename)
            
            if len(processed_files) == 1:
                return {'success': True, 'download_file': processed_files[0]}
            elif len(processed_files) > 1:
                zip_filename = '去公章.zip'
                zip_path = os.path.join(processed_dir, zip_filename)
                self._create_zip(processed_dir, processed_files, zip_path)
                return {'success': True, 'download_file': zip_filename}
            else:
                return {'error': '没有可处理的PDF文件'}
                
        except Exception as e:
            return {'error': f'处理失败: {str(e)}'}
    
    def convert_then_remove_header_seal(self, session_id, files):
        """转图片后去红头及公章"""
        try:
            processed_files = []
            processed_dir = os.path.join(self.processed_folder, session_id)
            os.makedirs(processed_dir, exist_ok=True)
            
            for file_info in files:
                if file_info['type'] == '文本型PDF':
                    input_path = os.path.join(self.upload_folder, session_id, file_info['filename'])
                    output_filename = f"去红头-{file_info['filename']}"
                    output_path = os.path.join(processed_dir, output_filename)
                    
                    # 先转换为图片型PDF，然后处理
                    temp_path = os.path.join(processed_dir, f"temp_{file_info['filename']}")
                    self._convert_text_to_image_pdf(input_path, temp_path)
                    self._remove_header_seal_scan_pdf(temp_path, output_path)
                    
                    # 删除临时文件
                    if os.path.exists(temp_path):
                        os.remove(temp_path)
                    
                    processed_files.append(output_filename)
            
            if len(processed_files) == 1:
                return {'success': True, 'download_file': processed_files[0]}
            elif len(processed_files) > 1:
                zip_filename = '去红头.zip'
                zip_path = os.path.join(processed_dir, zip_filename)
                self._create_zip(processed_dir, processed_files, zip_path)
                return {'success': True, 'download_file': zip_filename}
            else:
                return {'error': '没有可处理的文本型PDF文件'}
                
        except Exception as e:
            return {'error': f'处理失败: {str(e)}'}
    
    def convert_to_word(self, session_id, files):
        """转换为Word格式"""
        try:
            processed_files = []
            processed_dir = os.path.join(self.processed_folder, session_id)
            os.makedirs(processed_dir, exist_ok=True)
            
            for file_info in files:
                if file_info['type'] == '文本型PDF':
                    input_path = os.path.join(self.upload_folder, session_id, file_info['filename'])
                    output_filename = f"{os.path.splitext(file_info['filename'])[0]}.docx"
                    output_path = os.path.join(processed_dir, output_filename)
                    
                    self._pdf_to_word(input_path, output_path)
                    processed_files.append(output_filename)
            
            if len(processed_files) == 1:
                return {'success': True, 'download_file': processed_files[0]}
            elif len(processed_files) > 1:
                zip_filename = '转Word.zip'
                zip_path = os.path.join(processed_dir, zip_filename)
                self._create_zip(processed_dir, processed_files, zip_path)
                return {'success': True, 'download_file': zip_filename}
            else:
                return {'error': '没有可处理的文本型PDF文件'}
                
        except Exception as e:
            return {'error': f'处理失败: {str(e)}'}
    
    def _remove_header_seal_text_pdf(self, input_path, output_path):
        """处理文本型PDF：去红头和公章"""
        doc = fitz.open(input_path)
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            if page_num == 0:  # 第一页处理红头
                self._remove_header_from_page(page)
            
            # 所有页面处理公章
            self._remove_seal_from_page(page)
        
        # 使用增量保存和垃圾回收选项确保更改生效
        doc.save(output_path, garbage=4, deflate=True, clean=True)
        doc.close()
    
    def _remove_seal_text_pdf(self, input_path, output_path):
        """处理文本型PDF：仅去公章"""
        try:
            doc = fitz.open(input_path)
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                self._remove_seal_from_page(page)
            
            # 使用增量保存和垃圾回收选项确保更改生效
            doc.save(output_path, garbage=4, deflate=True, clean=True)
            doc.close()
            return True
        except Exception as e:
            print(f"[ERROR] 处理PDF时出错: {e}")
            return False
    
    def _remove_header_seal_scan_pdf(self, input_path, output_path):
        """处理扫描型PDF：去红头和公章 - 只处理有红头或公章的页面"""
        doc = fitz.open(input_path)
        new_doc = fitz.open()
        
        for page_num in range(len(doc)):
            print(f"\n[DEBUG] 处理第 {page_num + 1}/{len(doc)} 页...")
            page = doc[page_num]
            
            # 转换为图片，使用适中的分辨率以平衡质量和文件大小
            page_rect = page.rect
            mat = fitz.Matrix(1.5, 1.5)  # 降低分辨率从2.0到1.5
            pix = page.get_pixmap(matrix=mat)
            img_data = pix.tobytes("png")

            # 使用OpenCV处理图片
            nparr = np.frombuffer(img_data, np.uint8)
            img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)

            # 区分首页和非首页处理
            is_first_page = (page_num == 0)
            img = self._remove_all_red_from_image(img, is_first_page=is_first_page)

            # 转换回PDF页面，使用JPEG格式和压缩以减少文件大小
            encode_param = [int(cv2.IMWRITE_JPEG_QUALITY), 85]  # JPEG质量85%
            _, img_encoded = cv2.imencode('.jpg', img, encode_param)
            img_bytes = img_encoded.tobytes()

            # 创建新页面：使用原始页面的point尺寸，而非渲染像素尺寸，避免页面被物理放大
            img_rect = fitz.Rect(0, 0, page_rect.width, page_rect.height)
            new_page = new_doc.new_page(width=page_rect.width, height=page_rect.height)
            new_page.insert_image(img_rect, stream=img_bytes)

        print(f"\n[DEBUG] 所有页面处理完成")
        new_doc.save(output_path)
        new_doc.close()
        doc.close()

    def _remove_seal_scan_pdf(self, input_path, output_path):
        """处理扫描型PDF：仅去公章 - 首页上部63%区域不处理，只处理有公章的页面"""
        doc = fitz.open(input_path)
        new_doc = fitz.open()
        
        for page_num in range(len(doc)):
            print(f"\n[DEBUG] 处理第 {page_num + 1}/{len(doc)} 页...")
            page = doc[page_num]
            
            # 转换为图片，使用适中的分辨率以平衡质量和文件大小
            mat = fitz.Matrix(1.5, 1.5)  # 降低分辨率从2.0到1.5
            pix = page.get_pixmap(matrix=mat)
            img_data = pix.tobytes("png")
            
            # 使用OpenCV处理图片
            nparr = np.frombuffer(img_data, np.uint8)
            img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
            
            # 将红色改为白色，首页上部63%区域保护
            if page_num == 0:
                print("[DEBUG] 首页：保护上部63%区域（红头保护区）")
                # 首页：上部63%区域不处理，仅处理下部37%区域
                img = self._remove_red_with_header_protection(img, protect_ratio=0.63)
            else:
                print(f"[DEBUG] 第{page_num + 1}页：全页检测公章")
                # 非首页：全页检测公章并处理
                img = self._remove_all_red_from_image(img, is_first_page=False)

            # 转换回PDF页面，使用JPEG格式和压缩以减少文件大小
            encode_param = [int(cv2.IMWRITE_JPEG_QUALITY), 85]  # JPEG质量85%
            _, img_encoded = cv2.imencode('.jpg', img, encode_param)
            img_bytes = img_encoded.tobytes()

            # 创建新页面：使用原始页面的 point 尺寸，而非渲染出的像素尺寸，
            # 避免图片型PDF被按渲染倍率错误放大，导致后续水印字号相对变小
            page_rect = page.rect
            img_rect = fitz.Rect(0, 0, page_rect.width, page_rect.height)
            new_page = new_doc.new_page(width=page_rect.width, height=page_rect.height)
            new_page.insert_image(img_rect, stream=img_bytes)

        print(f"\n[DEBUG] 所有页面处理完成")
        new_doc.save(output_path)
        new_doc.close()
        doc.close()

    def _remove_header_from_page(self, page):
        """从文本型PDF页面中直接删除红头文字 - 通过修改内容流真正删除"""
        print(f"[DEBUG] 开始删除红头文字，页面尺寸: {page.rect}")
        
        # 动态红头删改（修正坐标方向，优先以红字簇确定切割线；否则退回正文顶部）
        page_width = page.rect.width
        page_height = page.rect.height

        header_bottom = None   # 红头文本的底边最大 y1（基于过滤后的红字）
        body_top = None        # 正文文本的顶边最小 y0（非红）
        spans_total = 0
        red_spans = 0
        nonred_spans = 0

        red_candidates = []  # 收集红字 (y0,y1) 候选
        try:
            text_dict = page.get_text("dict")
            for block in text_dict.get("blocks", []):
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        bbox = span.get("bbox")
                        if not bbox or len(bbox) < 4:
                            continue
                        x0, y0, x1, y1 = bbox  # (0,0) 左上，y 向下增大
                        text = (span.get("text") or "").strip()
                        if not text:
                            continue
                        spans_total += 1
                        color = span.get("color", 0)
                        if self._is_red_text(color):
                            red_spans += 1
                            # 先收集，后续基于 body_top 过滤
                            if y0 < page_height * 0.65:  # 只考虑上 65% 的红字
                                red_candidates.append((y0, y1))
                        else:
                            nonred_spans += 1
                            body_top = y0 if body_top is None else min(body_top, y0)
        except Exception as e:
            print(f"[DEBUG] 扫描文字失败: {e}")

        # 基于正文顶边过滤红字候选，计算 header_bottom
        if red_candidates:
            filtered = []
            if body_top is not None:
                margin = 1.0
                for y0, y1 in red_candidates:
                    if y1 <= max(0.0, body_top - margin):
                        filtered.append((y0, y1))
            else:
                filtered = red_candidates
            if filtered:
                header_bottom = max(y1 for _, y1 in filtered)
            else:
                # 若全部红字都在正文区域内（异常），不使用红字确定 header_bottom
                header_bottom = None

        # 计算 y_cut，双向约束不侵入正文
        pad_red = 2.0   # 红头到切割线缓冲
        pad_body = 12.0  # 切割线到正文缓冲（加大保护）
        y_cut = None
        if header_bottom is not None and body_top is not None:
            # 不越过正文顶部
            y_cut = min(header_bottom + pad_red, max(0.0, body_top - pad_body))
            y_cut = max(0.0, min(page_height, y_cut))
            print(f"[DEBUG] 红字+正文约束：header_bottom={header_bottom:.1f}, body_top={body_top:.1f} -> y_cut={y_cut:.1f}")
        elif header_bottom is not None:
            # 没有正文定位时，加上顶部上限保护（最多到 30% 页高）
            y_cut = min(header_bottom + pad_red, page_height * 0.30)
            y_cut = max(0.0, min(page_height, y_cut))
            print(f"[DEBUG] 仅红字约束（含上限）：header_bottom={header_bottom:.1f} -> y_cut={y_cut:.1f}")
        elif body_top is not None:
            y_cut = max(0.0, min(page_height, body_top - pad_body))
            print(f"[DEBUG] 仅正文约束：body_top={body_top:.1f} -> y_cut={y_cut:.1f}")
        else:
            top_ratio_fallback = 0.20
            y_cut = page_height * top_ratio_fallback
            print(f"[DEBUG] 兜底 20% 顶带 -> y_cut={y_cut:.1f}")

        # 安全保护：再次扫描正文，防止 y_cut 侵入黑色标题
        try:
            guard_top = None
            text_dict_guard = page.get_text("dict")
            for block in text_dict_guard.get("blocks", []):
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        text = (span.get("text") or "").strip()
                        if not text:
                            continue
                        # 只用更像正文的 span 做保护：字体大小≥7，长度≥1，非红（放宽以更稳妥保护标题）
                        if span.get("size", 0) >= 7 and len(text) >= 1 and not self._is_red_text(span.get("color", 0)):
                            bbox = span.get("bbox")
                            if not bbox or len(bbox) < 4:
                                continue
                            y0 = bbox[1]
                            guard_top = y0 if guard_top is None else min(guard_top, y0)
            if guard_top is not None:
                guard_margin = 18.0
                new_y_cut = min(y_cut, max(0.0, guard_top - guard_margin))
                if new_y_cut != y_cut:
                    print(f"[DEBUG] 正文保护收缩：guard_top={guard_top:.1f}, y_cut {y_cut:.1f} -> {new_y_cut:.1f}")
                y_cut = new_y_cut
        except Exception as e:
            print(f"[DEBUG] 正文保护检查失败: {e}")

        # 额外硬性保护：如果已识别正文顶部，强制 y_cut 不得超过 body_top - 14
        try:
            if body_top is not None:
                forced_y_cut = max(0.0, body_top - 14.0)
                if y_cut > forced_y_cut:
                    print(f"[DEBUG] 硬性保护：y_cut {y_cut:.1f} -> {forced_y_cut:.1f} (基于 body_top)")
                    y_cut = forced_y_cut
        except Exception:
            pass

        # 当检测到红字或明确的 header_bottom 时，使用“带状删改 + 交集矢量/图片补充”
        if red_spans > 0 or header_bottom is not None:
            # 统一对 [0, y_cut] 做删改（文本/矢量/图片）
            redact_rect = fitz.Rect(0, 0, page_width, y_cut)
            try:
                page.add_redact_annot(redact_rect, text="", fill=(1, 1, 1))
            except Exception as e:
                print(f"[DEBUG] 添加整体删改失败: {e}")

            # 补充：在删改带内的红色 span 再加一次删改（提高鲁棒）
            try:
                text_dict2 = page.get_text("dict")
                for block in text_dict2.get("blocks", []):
                    if block.get("type") != 0:
                        continue
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            bbox = span.get("bbox")
                            if not bbox:
                                continue
                            r = fitz.Rect(*bbox)
                            if r.y1 <= redact_rect.y1 + 0.5 and self._is_red_text(span.get("color", 0)):
                                try:
                                    page.add_redact_annot(r, text="", fill=(1, 1, 1))
                                except Exception:
                                    pass
            except Exception as e:
                print(f"[DEBUG] 补充红字删改失败: {e}")

            # 补充：删改带内相交的矢量与图片（仅交集）
            try:
                # 矢量
                for d in page.get_drawings():
                    rect = d.get("rect")
                    if rect is None:
                        continue
                    if isinstance(rect, (list, tuple)):
                        rect = fitz.Rect(*rect)
                    if rect:
                        inter = rect & redact_rect
                        if inter and inter.get_area() > 0:
                            try:
                                page.add_redact_annot(inter, text="", fill=(1,1,1))
                            except Exception:
                                pass
                # 图片
                for img in page.get_images(full=True):
                    xref = img[0]
                    try:
                        rects = page.get_image_rects(xref)
                    except Exception:
                        rects = []
                    for r in rects:
                        inter = r & redact_rect
                        if inter and inter.get_area() > 0:
                            try:
                                page.add_redact_annot(inter, text="", fill=(1,1,1))
                            except Exception:
                                pass
            except Exception as e:
                print(f"[DEBUG] 顶部矢量/图片删改补充失败: {e}")
        else:
            # 否则（无红字可用）：关闭带状删改，改为“仅删红色矢量/红色图像”，且不越过正文顶部
            limit_y = None
            if body_top is not None:
                limit_y = max(0.0, body_top - 14.0)
            else:
                limit_y = page_height * 0.18  # 兜底：只看顶端18%
            limit_band = fitz.Rect(0, 0, page_width, limit_y)
            print(f"[DEBUG] 无红字，采用定点删改：limit_y={limit_y:.1f}")

            # 仅删“红色”矢量（填充或描边为红），并裁剪到 limit_band 交集
            try:
                drawings = page.get_drawings()
                for d in drawings:
                    rect = d.get("rect")
                    if rect is None:
                        continue
                    if isinstance(rect, (list, tuple)):
                        rect = fitz.Rect(*rect)
                    # 检查颜色是否偏红（fill 或 stroke）
                    fill = d.get("fill")
                    stroke = d.get("color") or d.get("stroke")
                    def is_rgb_red(c):
                        try:
                            if isinstance(c, (list, tuple)) and len(c) >= 3:
                                r, g, b = c[0], c[1], c[2]
                                # 颜色值通常是0-1浮点
                                return (r is not None and r > 0.6) and (r > (g or 0) * 1.4) and (r > (b or 0) * 1.4)
                        except:
                            return False
                        return False
                    if rect and (is_rgb_red(fill) or is_rgb_red(stroke)):
                        inter = rect & limit_band
                        if inter and inter.get_area() > 0:
                            try:
                                page.add_redact_annot(inter, text="", fill=(1,1,1))
                            except Exception:
                                pass
            except Exception as e:
                print(f"[DEBUG] 红色矢量定点删改失败: {e}")

            # 仅删“红色”图片（位置在 limit_band 内），并裁剪到交集
            try:
                for img in page.get_images(full=True):
                    xref = img[0]
                    # 判断图片是否红色
                    is_red_img = False
                    try:
                        pix = fitz.Pixmap(page.parent, xref)
                        if pix.n - pix.alpha < 4:
                            img_bytes = pix.tobytes("png")
                            is_red_img = self._is_red_seal_image(img_bytes)
                        pix = None
                    except Exception:
                        is_red_img = False
                    if not is_red_img:
                        continue
                    # 位置与交集
                    try:
                        rects = page.get_image_rects(xref)
                    except Exception:
                        rects = []
                    for r in rects:
                        inter = r & limit_band
                        if inter and inter.get_area() > 0:
                            try:
                                page.add_redact_annot(inter, text="", fill=(1,1,1))
                            except Exception:
                                pass
            except Exception as e:
                print(f"[DEBUG] 红色图片定点删改失败: {e}")

        # 应用删改
        try:
            page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_REMOVE)
            print(f"[DEBUG] 红头删改完成（红字优先），red_spans={red_spans}, nonred_spans={nonred_spans}, y_cut={y_cut:.1f}")
            return
        except Exception as e:
            print(f"[DEBUG] 应用删改失败: {e}")
            return
        
        # 1. 删除文本形式的红头
        blocks = page.get_text("dict")["blocks"]
        removed_count = 0
        
        print(f"[DEBUG] 页面中共有 {len(blocks)} 个文本块")
        
        for block_idx, block in enumerate(blocks):
            if "lines" in block:
                for line_idx, line in enumerate(block["lines"]):
                    for span_idx, span in enumerate(line["spans"]):
                        bbox = span["bbox"]
                        text_content = span.get("text", "").strip()
                        
                        # 检查是否在红头区域（页面上部）
                        # bbox[1]是文本的y坐标（底部），bbox[3]是顶部
                        text_y = bbox[1]
                        if text_y > header_area_threshold:
                            print(f"[DEBUG] 红头区域文字 [{block_idx}.{line_idx}.{span_idx}]: '{text_content}' @ y={bbox[1]:.1f}")
                            
                            # 检查是否为红色文字
                            color_value = span.get("color", 0)
                            print(f"[DEBUG]   -> 颜色值: {color_value}")
                            is_red = self._is_red_text(color_value)
                            
                            if is_red:
                                if text_content:
                                    try:
                                        # 直接删除文本内容
                                        text_instances = page.search_for(text_content)
                                        print(f"[DEBUG]   -> 找到 {len(text_instances)} 个匹配实例")
                                        for inst in text_instances:
                                            # 检查实例是否在红头区域
                                            if inst.y0 < header_area:
                                                # 使用更精确的删除方法
                                                page.add_redact_annot(inst, text="", fill=(1, 1, 1))
                                                removed_count += 1
                                                print(f"[DEBUG]   -> 标记删除红头文字: '{text_content}' 位置: {inst}")
                                    except Exception as e:
                                        print(f"[DEBUG]   -> 删除文字时出错: {e}")
                            else:
                                print(f"[DEBUG]   -> 不是红色文字，跳过")
        
        # 应用所有删除操作
        if removed_count > 0:
            page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_REMOVE)
            print(f"[DEBUG] 成功删除 {removed_count} 个红头文字实例")
        
        # 2. 删除红色矢量图形（如红色线条、形状等）
        print(f"[DEBUG] 开始检测红色矢量图形")
        try:
            drawings = page.get_drawings()
            print(f"[DEBUG] 页面中发现 {len(drawings)} 个绘制对象")
            
            red_drawings_removed = 0
            for i, drawing in enumerate(drawings):
                # 打印绘制对象的详细信息用于调试
                if i < 5:  # 只打印前5个避免输出过多
                    print(f"[DEBUG] 绘制对象 {i}: {drawing}")
                
                # 检查绘制对象是否为红色且在红头区域
                if 'fill' in drawing and drawing['fill'] is not None:
                    fill_color = drawing['fill']
                    
                    # 检查是否为红色（放宽标准）
                    is_red_drawing = False
                    if isinstance(fill_color, (list, tuple)) and len(fill_color) >= 3:
                        r, g, b = fill_color[0], fill_color[1], fill_color[2]
                        # 放宽红色判断：红色分量 > 0.5 且明显高于绿蓝分量
                        is_red_drawing = r > 0.5 and r > g * 1.5 and r > b * 1.5
                        print(f"[DEBUG] 绘制对象 {i} 填充颜色: RGB({r:.2f}, {g:.2f}, {b:.2f}), 是否红色: {is_red_drawing}")
                    
                    if is_red_drawing:
                        # 获取绘制对象的边界矩形
                        if 'rect' in drawing:
                            rect = drawing['rect']
                            # 检查是否在红头区域
                            if rect.y0 < header_area:
                                print(f"[DEBUG] 发现红头区域红色绘制对象: {rect}")
                                try:
                                    page.add_redact_annot(rect, text="", fill=(1, 1, 1))
                                    red_drawings_removed += 1
                                    print(f"[DEBUG] 标记删除红色绘制对象: {rect}")
                                except Exception as e:
                                    print(f"[DEBUG] 标记删除红色绘制对象失败: {e}")
            
            # 应用绘制对象删除
            if red_drawings_removed > 0:
                page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_REMOVE)
                print(f"[DEBUG] 成功删除 {red_drawings_removed} 个红色绘制对象")
            else:
                print(f"[DEBUG] 没有找到需要删除的红色绘制对象")
                
        except Exception as e:
            print(f"[DEBUG] 处理红色绘制对象时出错: {e}")
        
        print(f"[DEBUG] 红头文字删除完成")
    
    def _remove_red_elements_from_stream(self, page, header_area):
        """通过编辑内容流删除红头区域的红色元素"""
        print("[DEBUG] 尝试通过内容流删除红头元素")
        
        try:
            # 读取页面内容流
            content_stream = page.read_contents()
            if not content_stream:
                print("[DEBUG] 无法读取页面内容流")
                return False
            
            content_str = content_stream.decode('latin-1', errors='ignore')
            original_content = content_str
            
            import re
            
            # 策略：找到红色指令后，删除直到下一个颜色改变或状态保存/恢复
            # 这样可以完全移除红色元素，而不是只改变颜色
            
            # 匹配红色设置到下一个关键指令之间的所有内容
            # 包括文本显示(Tj, TJ, ')、路径绘制(S, s, f, F, B)等
            red_block_pattern = r'(1(?:\.0*)?\s+0(?:\.0*)?\s+0(?:\.0*)?\s+(?:rg|RG)).*?(?=(?:\d+\.?\d*\s+){2,3}(?:rg|RG|sc|SC|scn|SCN|g|G|k|K)|q|Q|BT|ET|$)'
            
            matches = list(re.finditer(red_block_pattern, content_str, re.DOTALL))
            
            if matches:
                print(f"[DEBUG] 找到 {len(matches)} 个红色内容块")
                modifications_made = 0
                modified_content = content_str
                
                # 从后向前替换，避免位置偏移
                for match in reversed(matches):
                    start, end = match.span()
                    red_block = modified_content[start:end]
                    
                    # 检查这个块是否包含文本或图形绘制指令
                    has_content = bool(re.search(r'(Tj|TJ|\'|"|Td|TD|Tm|T\*|m|l|c|v|y|h|re|S|s|f|F\*?|B|b|W)', red_block))
                    
                    if has_content:
                        print(f"[DEBUG] 删除红色内容块 (长度: {len(red_block)})")
                        # 完全删除这个红色块
                        comment = f"% removed red block"
                        modified_content = modified_content[:start] + comment + modified_content[end:]
                        modifications_made += 1
                    else:
                        print(f"[DEBUG] 跳过空红色块")
                
                if modifications_made > 0:
                    content_str = modified_content
                else:
                    print("[DEBUG] 未找到需要删除的红色内容块")
                    # 如果上面的方法没找到，尝试简单的颜色替换
                    red_color_patterns = [
                        r'1(?:\.0+)?\s+0(?:\.0+)?\s+0(?:\.0+)?\s+rg',  # 红色填充
                        r'1(?:\.0+)?\s+0(?:\.0+)?\s+0(?:\.0+)?\s+RG',  # 红色描边
                    ]
                    
                    modifications_made = 0
                    modified_content = content_str
                    
                    for pattern in red_color_patterns:
                        matches_color = list(re.finditer(pattern, modified_content))
                        if matches_color:
                            print(f"[DEBUG] 找到 {len(matches_color)} 个红色指令: {pattern}")
                            # 从后向前替换为白色
                            for match in reversed(matches_color):
                                start, end = match.span()
                                original_cmd = modified_content[start:end]
                                if 'rg' in original_cmd.lower():
                                    replacement = '1 1 1 rg'
                                else:
                                    replacement = '1 1 1 RG'
                                modified_content = modified_content[:start] + replacement + modified_content[end:]
                                modifications_made += 1
                    
                    content_str = modified_content
            else:
                print("[DEBUG] 未找到红色内容块，尝试简单颜色替换")
                # 简单的颜色替换方法
                red_color_patterns = [
                    r'1(?:\.0+)?\s+0(?:\.0+)?\s+0(?:\.0+)?\s+rg',
                    r'1(?:\.0+)?\s+0(?:\.0+)?\s+0(?:\.0+)?\s+RG',
                ]
                
                modifications_made = 0
                modified_content = content_str
                
                for pattern in red_color_patterns:
                    matches_simple = list(re.finditer(pattern, modified_content))
                    if matches_simple:
                        print(f"[DEBUG] 找到 {len(matches_simple)} 个红色指令")
                        for match in reversed(matches_simple):
                            start, end = match.span()
                            modified_content = modified_content[:start] + '1 1 1 rg' + modified_content[end:]
                            modifications_made += 1
                
                content_str = modified_content
            
            # 检查是否有修改
            if content_str != original_content:
                print(f"[DEBUG] 内容流已修改，准备更新")
                try:
                    # 清理页面内容
                    page.clean_contents()
                    
                    # 获取内容流xref
                    contents_xref = page.get_contents()
                    if isinstance(contents_xref, list):
                        contents_xref = contents_xref[0] if contents_xref else None
                    
                    if contents_xref:
                        doc = page.parent
                        new_content_bytes = content_str.encode('latin-1', errors='ignore')
                        doc.update_stream(contents_xref, new_content_bytes)
                        print(f"[DEBUG] 成功更新内容流，删除红头元素")
                        return True
                    else:
                        print(f"[DEBUG] 无法获取内容流xref")
                        return False
                except Exception as e:
                    print(f"[DEBUG] 更新内容流失败: {e}")
                    import traceback
                    traceback.print_exc()
                    return False
            else:
                print("[DEBUG] 未找到需要修改的红色元素")
                return False
                
        except Exception as e:
            print(f"[DEBUG] 内容流编辑失败: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def _remove_red_vector_elements(self, page, header_area):
        """删除页面中红头区域的红色矢量元素"""
        try:
            # 获取页面内容流
            content = page.get_contents()
            if not content:
                return
            
            # 创建新页面来重构内容，过滤掉红色元素
            page_rect = page.rect
            
            # 转换页面为图像进行分析
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat)
            img_data = pix.tobytes("png")
            
            # 使用OpenCV检测红色区域
            nparr = np.frombuffer(img_data, np.uint8)
            img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
            
            if img is not None:
                # 只处理页面上部区域
                header_height = int(img.shape[0] * (header_area / page_rect.height))
                header_img = img[:header_height, :]
                
                # 检测红色区域
                hsv = cv2.cvtColor(header_img, cv2.COLOR_BGR2HSV)
                
                # 红色范围
                lower_red1 = np.array([0, 50, 50])
                upper_red1 = np.array([10, 255, 255])
                lower_red2 = np.array([170, 50, 50])
                upper_red2 = np.array([180, 255, 255])
                
                mask1 = cv2.inRange(hsv, lower_red1, upper_red1)
                mask2 = cv2.inRange(hsv, lower_red2, upper_red2)
                red_mask = mask1 + mask2
                
                # 查找轮廓
                contours, _ = cv2.findContours(red_mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
                
                for contour in contours:
                    area = cv2.contourArea(contour)
                    if area > 100:  # 最小面积阈值
                        # 获取轮廓的外接矩形
                        x, y, w, h = cv2.boundingRect(contour)
                        
                        # 转换回PDF坐标系
                        scale_x = page_rect.width / header_img.shape[1]
                        scale_y = header_area / header_img.shape[0]
                        
                        pdf_x = x * scale_x
                        pdf_y = y * scale_y
                        pdf_w = w * scale_x
                        pdf_h = h * scale_y
                        
                        # 使用redact annotation删除该区域
                        rect = fitz.Rect(pdf_x, pdf_y, pdf_x + pdf_w, pdf_y + pdf_h)
                        try:
                            page.add_redact_annot(rect)
                        except Exception:
                            pass
                
                # 应用删除操作
                try:
                    page.apply_redactions()
                except Exception:
                    pass
            
            pix = None
        except Exception:
            pass
    
    def _remove_seal_from_page(self, page):
        """从文本型PDF页面中删除公章 - 使用多种方法按优先级尝试"""
        print(f"[DEBUG] 开始删除公章，页面尺寸: {page.rect}")
        
        # 方法0：删除注释和印章（Annotations/Stamps）
        success0 = self._method0_remove_annotations(page)
        if success0:
            print("[DEBUG] 方法零成功删除注释形式的公章")
        
        # 方法一：PyMuPDF内置方法删除图像对象
        success1 = self._method1_pymupdf_image_deletion(page)
        if success1:
            print("[DEBUG] 方法一成功删除图像公章")
        
        # 方法1.5：删除矢量绘制的公章（页面中下部）
        success15 = self._remove_vector_seal_in_bottom(page)
        if success15:
            print("[DEBUG] 方法1.5成功删除矢量公章")
        
        # 汇总结果
        if success0 or success1 or success15:
            print(f"[DEBUG] 公章删除完成：方法0={success0}, 方法1={success1}, 方法1.5={success15}")
        else:
            print("[DEBUG] 所有方法均未检测到公章或删除失败")
    
    def _remove_annots_from_page_object(self, page):
        """直接从页面对象中删除注释（包括签名字段）"""
        print("[DEBUG] 尝试从页面对象中删除注释")
        
        try:
            doc = page.parent
            page_xref = page.xref
            
            # 获取当前页面的页码（从0开始）
            page_num = page.number
            print(f"[DEBUG] 当前页面页码: {page_num + 1}")
            
            # 读取页面对象
            page_obj = doc.xref_object(page_xref)
            
            # 检查是否有Annots键
            if '/Annots' not in page_obj:
                print("[DEBUG] 页面对象中没有/Annots键")
                return 0
            
            # 提取Annots数组中的xref
            import re
            annots_match = re.search(r'/Annots\s*\[\s*([^\]]+)\]', page_obj)
            if not annots_match:
                print("[DEBUG] 无法解析/Annots数组")
                return 0
            
            annots_content = annots_match.group(1)
            annot_xrefs = re.findall(r'(\d+)\s+0\s+R', annots_content)
            
            if not annot_xrefs:
                print("[DEBUG] /Annots数组为空")
                return 0
            
            print(f"[DEBUG] 页面对象中有 {len(annot_xrefs)} 个注释xref: {annot_xrefs}")
            
            # 检查是否有签名字段或公章（排除页面上部的注释）
            has_seal_in_middle_or_bottom = False
            page_height = page.rect.height
            
            # 定义页面上部区域（红头文字区域）
            # 红头只会出现在首页，非首页不需要保护红头区域
            # PDF坐标系：y=0在底部，y=page_height在顶部
            # 首页上部区域：y > page_height * 0.63（距底部大于63%），与扫描型PDF保持一致
            upper_area_y_threshold = page_height * 0.63
            is_first_page = (page_num == 0)
            
            for xref_str in annot_xrefs:
                xref = int(xref_str)
                
                try:
                    annot_obj = doc.xref_object(xref)
                    
                    # 检查是否是签名字段或Widget
                    is_signature = '/FT /Sig' in annot_obj
                    is_widget = '/Subtype /Widget' in annot_obj
                    has_stamp_keyword = any(kw in annot_obj for kw in ['Seal', 'Stamp', 'VISIBLE', '印章', '公章', 'Sign'])
                    
                    # 检查位置
                    rect_match = re.search(r'/Rect\s*\[\s*([\d\.\s]+)\]', annot_obj)
                    in_upper_area = False
                    in_middle_or_bottom = False
                    
                    if rect_match:
                        coords = rect_match.group(1).split()
                        if len(coords) >= 4:
                            x0, y0, x1, y1 = [float(c) for c in coords[:4]]
                            # y1是注释的顶部，y0是底部
                            # 判断注释的中心位置
                            y_center = (y0 + y1) / 2
                            from_bottom = page_height - y1
                            from_top = page_height - y_center
                            
                            # 判断是否在上部（红头区域）
                            # 红头只在首页存在，非首页的所有注释都可以删除
                            if is_first_page and y_center > upper_area_y_threshold:
                                in_upper_area = True
                                print(f"[DEBUG] 注释xref={xref}在首页上部（y中心={y_center:.1f}, 距顶部={from_top:.1f}），跳过")
                            else:
                                in_middle_or_bottom = True
                                if is_first_page:
                                    print(f"[DEBUG] 注释xref={xref}在首页中下部（y中心={y_center:.1f}, 距底部={from_bottom:.1f}）")
                                else:
                                    print(f"[DEBUG] 注释xref={xref}在第{page_num+1}页（y中心={y_center:.1f}），非首页直接可删除")
                    
                    # 只删除中下部的签名字段或印章
                    if (is_signature or (is_widget and has_stamp_keyword)) and in_middle_or_bottom:
                        print(f"[DEBUG] xref={xref}是中下部的签名字段/印章")
                        has_seal_in_middle_or_bottom = True
                        break
                    elif (is_signature or (is_widget and has_stamp_keyword)) and in_upper_area:
                        print(f"[DEBUG] xref={xref}在上部区域，不作为公章删除")
                
                except Exception as e:
                    print(f"[DEBUG] 检查注释xref={xref}失败: {e}")
                    continue
            
            # 如果发现中下部的签名字段或公章，彻底删除
            if has_seal_in_middle_or_bottom:
                print(f"[DEBUG] 发现中下部的签名字段/公章，开始彻底删除")
                
                # 步骤1：遍历所有中下部的签名字段，清除其外观流
                for xref_str in annot_xrefs:
                    xref = int(xref_str)
                    try:
                        annot_obj = doc.xref_object(xref)
                        is_signature = '/FT /Sig' in annot_obj
                        is_widget = '/Subtype /Widget' in annot_obj
                        has_stamp_keyword = any(kw in annot_obj for kw in ['Seal', 'Stamp', 'VISIBLE', '印章', '公章', 'Sign'])
                        
                        # 检查位置
                        rect_match = re.search(r'/Rect\s*\[\s*([\d\.\s]+)\]', annot_obj)
                        if rect_match:
                            coords = rect_match.group(1).split()
                            if len(coords) >= 4:
                                x0, y0, x1, y1 = [float(c) for c in coords[:4]]
                                y_center = (y0 + y1) / 2
                                
                                # 如果是签名字段（首页需要判断位置，非首页直接处理）
                                if (is_signature or (is_widget and has_stamp_keyword)) and (not is_first_page or y_center <= upper_area_y_threshold):
                                    if not is_first_page:
                                        print(f"[DEBUG] 第{page_num+1}页的签名字段 xref={xref}，非首页直接删除")
                                    else:
                                        print(f"[DEBUG] 首页中下部的签名字段 xref={xref}")
                                    print(f"[DEBUG] 清除签名字段 xref={xref} 的外观流")
                                    
                                    # 方法1：提取并删除/AP引用的外观流对象
                                    try:
                                        ap_match = re.search(r'/AP\s*<<\s*/N\s+(\d+)\s+0\s+R', annot_obj)
                                        if ap_match:
                                            ap_xref = int(ap_match.group(1))
                                            print(f"[DEBUG] 发现/AP/N引用 xref={ap_xref}")
                                            
                                            # 尝试清空外观流对象的内容
                                            try:
                                                # 将外观流对象替换为空字典
                                                empty_obj = "<< /Length 0 >>\nstream\n\nendstream"
                                                doc.update_object(ap_xref, empty_obj)
                                                print(f"[DEBUG] 已清空外观流对象 xref={ap_xref}")
                                            except Exception as e:
                                                print(f"[DEBUG] 清空外观流对象失败: {e}")
                                    except Exception as e:
                                        print(f"[DEBUG] 提取/AP引用失败: {e}")
                                    
                                    # 方法2：重构整个签名字段对象，移除/AP、/V、/DR键
                                    try:
                                        # 读取原对象
                                        original_obj = doc.xref_object(xref)
                                        
                                        # 构建新对象：保留必要的键，移除/AP、/V、/DR
                                        new_obj_parts = []
                                        new_obj_parts.append("<<")
                                        
                                        # 提取并保留必要的键
                                        for key in ['/FT', '/T', '/F', '/Type', '/Subtype', '/Rect', '/P']:
                                            match = re.search(rf'{re.escape(key)}\s+([^\n/]+)', original_obj)
                                            if match:
                                                value = match.group(1).strip()
                                                # 移除尾部的额外空格或换行
                                                if value and not value.endswith('>>'):
                                                    new_obj_parts.append(f"  {key} {value}")
                                        
                                        new_obj_parts.append(">>")
                                        new_obj = "\n".join(new_obj_parts)
                                        
                                        print(f"[DEBUG] 重构签名字段对象 xref={xref}")
                                        print(f"[DEBUG] 新对象: {new_obj}")
                                        doc.update_object(xref, new_obj)
                                        print(f"[DEBUG] 已重构 xref={xref}，移除了/AP、/V、/DR键")
                                    except Exception as e:
                                        print(f"[DEBUG] 重构签名字段对象失败: {e}")
                                        import traceback
                                        traceback.print_exc()
                    except Exception as e:
                        print(f"[DEBUG] 处理签名字段 xref={xref} 失败: {e}")
                        continue
                
                # 步骤2：删除页面的/Annots键
                doc.xref_set_key(page_xref, "Annots", "null")
                print(f"[DEBUG] 已将/Annots设置为null")
                return 1
            else:
                print(f"[DEBUG] 未发现中下部的签名字段或公章")
                return 0
            
        except Exception as e:
            print(f"[DEBUG] 从页面对象删除注释失败: {e}")
            import traceback
            traceback.print_exc()
            return 0
    
    def _remove_stamps_from_widgets(self, page):
        """从表单字段（Widgets）中删除印章"""
        print("[DEBUG] 检查表单字段中的印章")
        
        try:
            widgets = page.widgets()
            if not widgets:
                print("[DEBUG] 页面中没有表单字段")
                return 0
            
            removed_count = 0
            for widget in widgets:
                try:
                    widget_type = widget.field_type
                    widget_name = widget.field_name
                    
                    print(f"[DEBUG] 发现表单字段: 类型={widget_type}, 名称={widget_name}")
                    
                    # 检查字段名称或类型
                    if any(keyword in str(widget_name).lower() for keyword in ['stamp', 'seal', '印章', '公章', 'sign']):
                        print(f"[DEBUG] 表单字段 '{widget_name}' 可能是印章，尝试删除")
                        # 尝试删除这个widget
                        try:
                            page.delete_widget(widget)
                            removed_count += 1
                            print(f"[DEBUG] 成功删除表单字段印章")
                        except Exception as e:
                            print(f"[DEBUG] 删除表单字段失败: {e}")
                except Exception as e:
                    print(f"[DEBUG] 处理表单字段时出错: {e}")
                    continue
            
            if removed_count > 0:
                print(f"[DEBUG] 从表单字段中删除了 {removed_count} 个印章")
            
            return removed_count
        except Exception as e:
            print(f"[DEBUG] 检查表单字段失败: {e}")
            return 0
    
    def _method0_remove_annotations(self, page):
        """方法零：删除注释和印章（Annotations/Stamps）以及表单字段"""
        print("[DEBUG] 尝试方法零：删除注释、印章和表单字段")
        
        try:
            # 首先尝试直接从页面对象中删除签名注释
            removed_from_page_obj = self._remove_annots_from_page_object(page)
            
            # 如果已经从页面对象删除了注释，直接返回成功
            if removed_from_page_obj > 0:
                print(f"[DEBUG] 从页面对象成功删除 {removed_from_page_obj} 个注释")
                return True
            
            # 然后处理表单字段中的印章
            removed_from_widgets = self._remove_stamps_from_widgets(page)
            if removed_from_widgets > 0:
                print(f"[DEBUG] 从表单字段成功删除 {removed_from_widgets} 个印章")
                return True
            
            annots = page.annots()
            if not annots:
                print("[DEBUG] 页面中没有可遍历的注释")
                return False
            
            removed_count = 0
            annot = page.first_annot
            
            while annot:
                next_annot = annot.next
                
                # 获取注释类型和信息
                annot_type = annot.type[0] if annot.type else -1
                annot_info = annot.info
                annot_name = annot_info.get('name', '') if isinstance(annot_info, dict) else ''
                
                print(f"[DEBUG] 发现注释: 类型={annot_type}, 名称={annot_name}, 信息={annot_info}")
                
                # 检查是否是印章类型的注释
                # Type 13 = Stamp (印章)
                # Type 0 = Text (文本)
                # Type 17 = Redact (涂改)
                # Type 24 = Widget (表单字段)
                should_remove = False
                
                if annot_type == 13:  # Stamp类型
                    print(f"[DEBUG] 发现印章注释")
                    should_remove = True
                elif annot_type == 17:  # Redact类型
                    print(f"[DEBUG] 发现涂改注释")
                    should_remove = True
                elif annot_type == 24:  # Widget类型（表单字段）
                    print(f"[DEBUG] 发现表单字段，检查是否为印章")
                    # 检查字段名称是否包含印章相关关键词
                    if any(keyword in annot_name.lower() for keyword in ['stamp', 'seal', '印章', '公章']):
                        print(f"[DEBUG] 表单字段名称包含印章关键词")
                        should_remove = True
                else:
                    # 检查注释的外观是否包含红色
                    try:
                        # 获取注释的外观流
                        ap = annot.get_pixmap()
                        if ap:
                            # 检查是否有红色像素
                            img_data = ap.tobytes("png")
                            if self._is_red_seal_image(img_data):
                                print(f"[DEBUG] 注释包含红色内容，可能是公章")
                                should_remove = True
                            ap = None
                    except Exception as e:
                        print(f"[DEBUG] 无法检查注释外观: {e}")
                
                if should_remove:
                    try:
                        page.delete_annot(annot)
                        removed_count += 1
                        print(f"[DEBUG] 删除注释: 类型={annot_type}")
                    except Exception as e:
                        print(f"[DEBUG] 删除注释失败: {e}")
                
                annot = next_annot
            
            if removed_count > 0:
                print(f"[DEBUG] 方法零成功删除 {removed_count} 个注释")
                return True
            else:
                print("[DEBUG] 方法零未删除任何注释")
                return False
                
        except Exception as e:
            print(f"[DEBUG] 方法零执行失败: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def _method1_pymupdf_image_deletion(self, page):
        """方法一：PyMuPDF内置方法删除图像对象"""
        print("[DEBUG] 尝试方法一：PyMuPDF内置方法删除图像")
        
        image_list = page.get_images()
        print(f"[DEBUG] 页面中发现 {len(image_list)} 个图片")
        
        images_removed = 0
        seal_images = []  # 记录所有红色公章
        
        # 第一步：检测所有红色公章（只检测中下部的）
        page_height = page.rect.height
        
        for img_index, img in enumerate(image_list):
            try:
                # 获取图片数据
                xref = img[0]
                pix = fitz.Pixmap(page.parent, xref)
                
                if pix.n - pix.alpha < 4:  # 确保是RGB图片
                    img_data = pix.tobytes("png")
                    
                    # 检查是否为红色图像
                    is_red = self._is_red_seal_image(img_data)
                    print(f"[DEBUG] 图片 {img_index} 是否为红色图像: {is_red}")
                    
                    if is_red:
                        # 检查图像位置，只处理中下部的
                        try:
                            img_rects = page.get_image_rects(xref)
                            if img_rects:
                                rect = img_rects[0]
                                y_center = (rect.y0 + rect.y1) / 2
                                from_bottom = page_height - rect.y1
                                
                                # 只删除中下部的红色图像
                                # from_bottom < 500 表示距底部较近（中下部）
                                # 上部的红色图像（如红头，from_bottom > 600）不删除
                                if from_bottom < 500:
                                    img_name = img[7]
                                    seal_images.append((img_index, img_name, xref))
                                    print(f"[DEBUG] 检测到中下部红色公章: 图片{img_index}, 名称={img_name}, y中心={y_center:.1f}, 距底部={from_bottom:.1f}")
                                else:
                                    print(f"[DEBUG] 图片{img_index}在上部（y中心={y_center:.1f}, 距底部={from_bottom:.1f}），跳过")
                        except Exception as e:
                            print(f"[DEBUG] 无法获取图片{img_index}位置: {e}")
                
                pix = None
            except Exception as e:
                print(f"[DEBUG] 处理图片 {img_index} 时出错: {e}")
                continue
        
        # 第二步：批量从内容流中删除所有公章
        if seal_images:
            print(f"\n[DEBUG] 共检测到 {len(seal_images)} 个红色公章，开始删除...")
            
            try:
                # 读取页面内容流
                content = page.read_contents()
                content_str = content.decode('latin-1', errors='ignore')
                modified_content = content_str
                
                import re
                
                doc = page.parent
                
                # 逐个删除每个公章
                for img_index, img_name, xref in seal_images:
                    print(f"[DEBUG] 删除公章 {img_name}...")
                    
                    # 步骤1: 替换图像数据为1x1透明图像
                    try:
                        # 创建1x1完全透明的PNG
                        transparent_png = b'\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x06\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\x0bIDATx\x9cc\x00\x01\x00\x00\x05\x00\x01\r\n-\xb4\x00\x00\x00\x00IEND\xaeB`\x82'
                        doc.update_stream(xref, transparent_png)
                        print(f"  步骤1: 图像数据已替换为透明")
                    except Exception as e:
                        print(f"  步骤1失败: {e}")
                    
                    # 步骤2: 从内容流中删除引用
                    # 模式1: 完整的 q...cm/ImXX Do Q
                    pattern1 = r'q\s+[\d\.\s\-]+cm\s*/' + re.escape(img_name) + r'\s+Do\s+Q'
                    temp_content = re.sub(pattern1, '', modified_content)
                    
                    if temp_content != modified_content:
                        print(f"  步骤2: 使用模式1删除内容流引用")
                        modified_content = temp_content
                        images_removed += 1
                    else:
                        # 模式2: 简单的 /ImXX Do
                        pattern2 = r'/' + re.escape(img_name) + r'\s+Do'
                        if re.search(pattern2, modified_content):
                            temp_content = re.sub(pattern2, '', modified_content)
                            if temp_content != modified_content:
                                print(f"  步骤2: 使用模式2删除内容流引用")
                                modified_content = temp_content
                                images_removed += 1
                        else:
                            print(f"  步骤2: 未找到内容流引用")
                
                # 如果有修改，更新内容流
                if modified_content != content_str:
                    page.clean_contents()
                    contents_xref = page.get_contents()
                    if isinstance(contents_xref, list):
                        contents_xref = contents_xref[0] if contents_xref else None
                    
                    if contents_xref:
                        doc = page.parent
                        new_content_bytes = modified_content.encode('latin-1', errors='ignore')
                        doc.update_stream(contents_xref, new_content_bytes)
                        print(f"[DEBUG] 内容流已更新，删除了 {images_removed} 个公章引用")
                        print(f"[DEBUG] 内容流大小: {len(content_str)} -> {len(modified_content)} 字节")
            
            except Exception as e:
                print(f"[DEBUG] 批量删除失败: {e}")
                import traceback
                traceback.print_exc()
        
        # 返回删除结果
        if images_removed > 0:
            print(f"[DEBUG] 方法一成功删除 {images_removed} 个公章图片")
            return True
        
        return False
    
    def _remove_vector_seal_in_bottom(self, page):
        """删除页面中下部的红色矢量公章"""
        print("[DEBUG] 检测并删除页面中下部的红色矢量公章")
        
        try:
            content = page.read_contents()
            if not content:
                return False
            
            content_str = content.decode('latin-1', errors='ignore')
            page_height = page.rect.height
            
            import re
            
            # 策略：找到页面中下部（y > page_height/2）的红色绘图块并删除
            # 1. 找到所有红色颜色设置指令
            # 2. 向后查找直到遇到下一个颜色指令或状态改变
            # 3. 分析该块中的坐标，判断是否在中下部
            # 4. 如果是，删除整个块
            
            # 查找所有红色指令及其位置
            red_pattern = r'(1(?:\.0+)?\s+0(?:\.0+)?\s+0(?:\.0+)?\s+(?:rg|RG))'
            red_matches = list(re.finditer(red_pattern, content_str))
            
            if not red_matches:
                print("[DEBUG] 未找到红色指令")
                return False
            
            print(f"[DEBUG] 找到 {len(red_matches)} 个红色指令")
            
            modified_content = content_str
            removed_count = 0
            
            # 从后向前处理，避免位置偏移
            for match in reversed(red_matches):
                red_start = match.start()
                red_end = match.end()
                
                # 查找这个红色块的结束位置（下一个颜色指令或Q）
                next_color_pattern = r'(?:\d+\.?\d*\s+){2,3}(?:rg|RG|sc|SC|scn|SCN|g|G|k|K)|[Qq]'
                next_match = re.search(next_color_pattern, content_str[red_end:])
                
                if next_match:
                    block_end = red_end + next_match.start()
                else:
                    # 如果找不到结束，取到内容流结尾或下500个字符
                    block_end = min(red_end + 500, len(content_str))
                
                red_block = content_str[red_start:block_end]
                
                # 分析这个块中的y坐标
                # 查找移动和绘图指令中的坐标: m, l, c, re等
                coord_pattern = r'([\d\.]+)\s+([\d\.]+)\s+m|' \
                               r'([\d\.]+)\s+([\d\.]+)\s+l|' \
                               r'([\d\.]+)\s+([\d\.]+)\s+([\d\.]+)\s+([\d\.]+)\s+re'
                
                coords = re.findall(coord_pattern, red_block)
                
                if coords:
                    # 提取所有y坐标
                    y_coords = []
                    for coord_match in coords:
                        # coord_match是一个元组，包含多个捕获组
                        for i in range(1, len(coord_match), 2):
                            if coord_match[i]:
                                try:
                                    y = float(coord_match[i])
                                    y_coords.append(y)
                                except:
                                    pass
                    
                    if y_coords:
                        avg_y = sum(y_coords) / len(y_coords)
                        from_bottom = page_height - avg_y
                        
                        # 判断是否在页面中下部（距离底部100-400之间，排除页眉页脚）
                        if 100 < from_bottom < 400:
                            print(f"[DEBUG] 发现中下部红色绘图: y={avg_y:.1f}, 距底部={from_bottom:.1f}")
                            # 删除整个红色块
                            modified_content = modified_content[:red_start] + \
                                             f"% removed bottom seal {removed_count}" + \
                                             modified_content[block_end:]
                            removed_count += 1
            
            if removed_count > 0:
                print(f"[DEBUG] 删除了 {removed_count} 个中下部红色块")
                # 更新内容流
                page.clean_contents()
                contents_xref = page.get_contents()
                if isinstance(contents_xref, list):
                    contents_xref = contents_xref[0] if contents_xref else None
                
                if contents_xref:
                    doc = page.parent
                    new_content_bytes = modified_content.encode('latin-1', errors='ignore')
                    doc.update_stream(contents_xref, new_content_bytes)
                    print(f"[DEBUG] 成功删除中下部矢量公章")
                    return True
            
            return False
            
        except Exception as e:
            print(f"[DEBUG] 删除矢量公章失败: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def _method2_content_stream_editing(self, page):
        """方法二：内容流编辑删除公章绘图指令"""
        print("[DEBUG] 尝试方法二：内容流编辑删除方法")
        
        try:
            # 获取页面内容流
            content_stream = page.read_contents()
            if not content_stream:
                print("[DEBUG] 无法读取页面内容流")
                return False
            
            content_str = content_stream.decode('latin-1', errors='ignore')
            original_content = content_str
            
            # 查找并替换红色绘制指令
            import re
            
            # 打印内容流的统计信息
            print(f"[DEBUG] 内容流长度: {len(content_str)} 字节")
            
            # 查找所有颜色指令
            all_colors = re.findall(r'(\d+\.?\d*)\s+(\d+\.?\d*)\s+(\d+\.?\d*)\s+(rg|RG|sc|SC)', content_str)
            if all_colors:
                print(f"[DEBUG] 发现 {len(all_colors)} 个颜色指令")
                # 分析前5个颜色指令
                for i, (r, g, b, cmd) in enumerate(all_colors[:5]):
                    print(f"[DEBUG] 颜色 {i}: RGB({r}, {g}, {b}) {cmd}")
            
            # 定义各种红色模式 - 更精确和更宽松的匹配
            red_patterns = [
                # 标准红色
                r'1(?:\.0+)?\s+0(?:\.0+)?\s+0(?:\.0+)?\s+rg',  # 红色填充
                r'1(?:\.0+)?\s+0(?:\.0+)?\s+0(?:\.0+)?\s+RG',  # 红色描边
                # 深红色
                r'0\.[89]\d*\s+0\.0\d*\s+0\.0\d*\s+rg',
                r'0\.[89]\d*\s+0\.0\d*\s+0\.0\d*\s+RG',
                # 更宽松的红色匹配（红色分量 > 0.7，其他 < 0.3）
                r'0\.[789]\d*\s+0\.[012]\d*\s+0\.[012]\d*\s+rg',
                r'0\.[789]\d*\s+0\.[012]\d*\s+0\.[012]\d*\s+RG',
            ]
            
            modifications_made = 0
            modified_content = content_str
            
            for pattern in red_patterns:
                matches = list(re.finditer(pattern, modified_content))
                if matches:
                    print(f"[DEBUG] 找到 {len(matches)} 个匹配: {pattern}")
                    # 从后向前替换，避免位置偏移问题
                    for match in reversed(matches):
                        # 将红色指令注释掉
                        start, end = match.span()
                        original_cmd = modified_content[start:end]
                        # 替换为空操作（保持PDF结构完整）
                        comment = f"% removed red: {original_cmd}"
                        modified_content = modified_content[:start] + comment + modified_content[end:]
                        modifications_made += 1
            
            # 如果有修改，更新页面内容
            if modifications_made > 0:
                print(f"[DEBUG] 共标记删除 {modifications_made} 个红色绘制指令")
                try:
                    # 清理页面内容流
                    page.clean_contents()
                    
                    # 获取内容流的xref
                    contents_xref = page.get_contents()
                    if isinstance(contents_xref, list):
                        # 如果有多个内容流，只更新第一个
                        contents_xref = contents_xref[0] if contents_xref else None
                    
                    if contents_xref:
                        # 直接更新文档中的stream对象
                        doc = page.parent
                        new_content_bytes = modified_content.encode('latin-1', errors='ignore')
                        
                        # 使用update_stream更新
                        doc.update_stream(contents_xref, new_content_bytes)
                        print(f"[DEBUG] 方法二成功更新页面内容流")
                        return True
                    else:
                        print(f"[DEBUG] 无法获取内容流xref")
                        return False
                        
                except Exception as e:
                    print(f"[DEBUG] 更新页面内容失败: {e}")
                    import traceback
                    traceback.print_exc()
                    return False
            else:
                print("[DEBUG] 方法二未发现需要修改的红色绘制指令")
                return False
                
        except Exception as e:
            print(f"[DEBUG] 方法二执行失败: {e}")
            return False
    
    def _method3_resource_dict_editing(self, page):
        """方法三：资源字典编辑删除公章图像引用"""
        print("[DEBUG] 尝试方法三：资源字典编辑删除方法")
        
        try:
            # 获取页面对象字典
            page_obj = page.get_contents()
            
            # 尝试获取页面资源
            try:
                # 方法1: 通过页面内容查找资源引用
                content_stream = page.read_contents()
                content_str = content_stream.decode('latin-1', errors='ignore')
                
                # 查找XObject引用
                import re
                xobject_refs = re.findall(r'/(\w+)\s+Do', content_str)
                
                if xobject_refs:
                    print(f"[DEBUG] 在内容流中发现XObject引用: {xobject_refs}")
                    
                    # 尝试删除这些引用
                    modifications_made = 0
                    for ref in xobject_refs:
                        # 检查引用前的颜色设置
                        ref_pattern = f'/{ref}\\s+Do'
                        matches = list(re.finditer(ref_pattern, content_str))
                        
                        for match in matches:
                            start_pos = max(0, match.start() - 200)
                            context = content_str[start_pos:match.start()]
                            
                            # 如果在上下文中发现红色设置，删除该引用
                            if re.search(r'1\.?0*\s+0\.?0*\s+0\.?0*\s+rg', context) or \
                               re.search(r'0\.[89]\d*\s+0\.[01]\d*\s+0\.[01]\d*\s+rg', context):
                                # 替换为注释
                                content_str = content_str.replace(match.group(), f'% removed {ref} Do')
                                modifications_made += 1
                                print(f"[DEBUG] 删除了XObject引用: {ref}")
                    
                    if modifications_made > 0:
                        # 更新页面内容
                        try:
                            new_content = content_str.encode('latin-1')
                            page.set_contents(new_content)
                            print(f"[DEBUG] 方法三成功修改了 {modifications_made} 个XObject引用")
                            return True
                        except Exception as e:
                            print(f"[DEBUG] 更新页面内容失败: {e}")
                            return False
                else:
                    print("[DEBUG] 未在内容流中发现XObject引用")
                    return False
                    
            except Exception as e:
                print(f"[DEBUG] 方法三执行失败: {e}")
                return False
            
            return False
                
        except Exception as e:
            print(f"[DEBUG] 方法三执行失败: {e}")
            return False
    
    def _method4_transparent_replacement(self, page):
        """方法四：透明图像替换方法"""
        print("[DEBUG] 尝试方法四：透明图像替换方法")
        
        try:
            image_list = page.get_images()
            replacements_made = 0
            
            for img_index, img in enumerate(image_list):
                try:
                    xref = img[0]
                    pix = fitz.Pixmap(page.parent, xref)
                    
                    if pix.n - pix.alpha < 4:  # RGB图像
                        img_data = pix.tobytes("png")
                        
                        if self._is_red_seal_image(img_data):
                            # 创建透明图像替换原图像
                            try:
                                # 获取原图像尺寸
                                width = pix.width
                                height = pix.height
                                
                                # 创建同尺寸的透明图像
                                transparent_pix = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, width, height), True)
                                transparent_pix.clear_with(255)  # 填充为白色
                                transparent_pix.set_alpha(0)     # 设置为完全透明
                                
                                # 将透明图像数据更新到原xref
                                transparent_data = transparent_pix.tobytes("png")
                                page.parent.update_stream(xref, transparent_data)
                                
                                replacements_made += 1
                                print(f"[DEBUG] 成功用透明图像替换了公章图像 {img_index}")
                                
                                transparent_pix = None
                                
                            except Exception as e:
                                print(f"[DEBUG] 透明图像替换失败: {e}")
                    
                    pix = None
                    
                except Exception as e:
                    print(f"[DEBUG] 处理图像 {img_index} 时出错: {e}")
                    continue
            
            if replacements_made > 0:
                print(f"[DEBUG] 方法四成功替换了 {replacements_made} 个公章图像")
                return True
            else:
                print("[DEBUG] 方法四未发现需要替换的图像")
                return False
                
        except Exception as e:
            print(f"[DEBUG] 方法四执行失败: {e}")
            return False
    
    def _remove_header_from_image(self, img):
        """从图片中移除红头"""
        height, width = img.shape[:2]
        header_height = int(height * 0.3)  # 上部30%区域
        
        # 在红头区域创建白色遮罩
        header_area = img[:header_height, :]
        
        # 检测红色区域
        hsv = cv2.cvtColor(header_area, cv2.COLOR_BGR2HSV)
        
        # 红色的HSV范围
        lower_red1 = np.array([0, 50, 50])
        upper_red1 = np.array([10, 255, 255])
        lower_red2 = np.array([170, 50, 50])
        upper_red2 = np.array([180, 255, 255])
        
        mask1 = cv2.inRange(hsv, lower_red1, upper_red1)
        mask2 = cv2.inRange(hsv, lower_red2, upper_red2)
        red_mask = mask1 + mask2
        
        # 用白色覆盖红色区域
        img[:header_height, :][red_mask > 0] = [255, 255, 255]
        
        return img
    
    def _remove_all_red_from_image(self, img, is_first_page=False):
        """将图片中所有红色改为白色，同时保护黑色文字
        
        Args:
            img: 输入图片
            is_first_page: 是否为首页（只有首页可能有红头）
        """
        img_height = img.shape[0]
        
        if is_first_page:
            # 首页：分两部分处理
            # 上部70%可能有红头，下部30%可能有公章
            header_boundary = int(img_height * 0.70)
            
            # 上部70%：检测红头并激进处理
            if header_boundary > 0:
                header_part = img[:header_boundary, :].copy()
                if self._has_red_header(header_part):
                    print("[DEBUG] 首页上部检测到红头，激进去除...")
                    header_part = self._remove_red_aggressive(header_part)
                    img[:header_boundary, :] = header_part
                else:
                    print("[DEBUG] 首页上部未检测到红头")
            
            # 下部30%：检测公章并保护文字处理
            if header_boundary < img_height:
                bottom_part = img[header_boundary:, :].copy()
                if self._has_red_seal(bottom_part):
                    print("[DEBUG] 首页下部检测到公章，保护文字去除...")
                    bottom_part = self._remove_red_protect_black_text(bottom_part)
                    img[header_boundary:, :] = bottom_part
                else:
                    print("[DEBUG] 首页下部未检测到公章")
        else:
            # 非首页：全页检测公章，使用保护文字的方法
            # 非首页不会有红头，只可能有公章
            if self._has_red_seal(img):
                print("[DEBUG] 非首页检测到公章，保护文字去除...")
                img = self._remove_red_protect_black_text(img)
            else:
                print("[DEBUG] 非首页未检测到公章，跳过处理")
        
        return img
    
    def _has_red_header(self, img):
        """检测图像上部是否有红头文字
        红头特征：位于页面上部，红色文字行，行间无黑色文字"""
        # 转换为HSV颜色空间
        hsv = cv2.cvtColor(img, cv2.COLOR_BGR2HSV)
        
        # 检测红色区域（使用非常宽松的阈值）
        lower_red1 = np.array([0, 20, 20])  # 降低阈值
        upper_red1 = np.array([25, 255, 255])  # 扩大范围
        lower_red2 = np.array([155, 20, 20])  # 扩大范围
        upper_red2 = np.array([180, 255, 255])
        
        mask1 = cv2.inRange(hsv, lower_red1, upper_red1)
        mask2 = cv2.inRange(hsv, lower_red2, upper_red2)
        red_mask = mask1 + mask2
        
        # 计算红色像素占比
        total_pixels = img.shape[0] * img.shape[1]
        red_pixels = np.sum(red_mask > 0)
        red_ratio = red_pixels / total_pixels
        
        # 降低阈值：红色占比超过0.1%就认为可能有红头
        if red_ratio > 0.001:
            print(f"[DEBUG] 红头检测：红色占比 {red_ratio:.2%}，检测到红头")
            return True
        
        print(f"[DEBUG] 红头检测：红色占比过低 {red_ratio:.2%}，无红头")
        return False
    
    def _has_red_seal(self, img):
        """检测图像中是否有红色圆形公章
        公章特征：红色圆形，直径3.5-5cm，可能与黑色文字重叠"""
        # 转换为HSV颜色空间
        hsv = cv2.cvtColor(img, cv2.COLOR_BGR2HSV)
        
        # 检测红色区域（使用宽松阈值）
        lower_red1 = np.array([0, 20, 20])  # 降低阈值
        upper_red1 = np.array([25, 255, 255])  # 扩大范围
        lower_red2 = np.array([155, 20, 20])  # 扩大范围
        upper_red2 = np.array([180, 255, 255])
        
        mask1 = cv2.inRange(hsv, lower_red1, upper_red1)
        mask2 = cv2.inRange(hsv, lower_red2, upper_red2)
        red_mask = mask1 + mask2
        
        # 计算红色像素占比
        total_pixels = img.shape[0] * img.shape[1]
        red_pixels = np.sum(red_mask > 0)
        red_ratio = red_pixels / total_pixels
        
        # 降低阈值：红色占比 > 0.1% 就检测
        if red_ratio < 0.001:
            print(f"[DEBUG] 公章检测：红色占比过低 {red_ratio:.2%}，无公章")
            return False
        
        # 形态学操作，连接断裂的红色区域
        kernel = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (5, 5))
        red_mask = cv2.morphologyEx(red_mask, cv2.MORPH_CLOSE, kernel, iterations=2)
        
        # 查找轮廓
        contours, _ = cv2.findContours(red_mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        
        if not contours:
            print("[DEBUG] 公章检测：未找到红色轮廓")
            return False
        
        # 检查是否有圆形轮廓（可能是公章）
        for contour in contours:
            area = cv2.contourArea(contour)
            
            # 公章面积范围：直径3.5-5cm
            # 在150dpi下，直径约200-300像素，面积 = π * r² ≈ 30000-70000
            # 放宽范围以适应不同分辨率
            if 3000 < area < 300000:  # 进一步放宽范围
                # 计算圆形度
                perimeter = cv2.arcLength(contour, True)
                if perimeter > 0:
                    circularity = 4 * np.pi * area / (perimeter * perimeter)
                    # 降低圆形度要求
                    if circularity > 0.3:
                        print(f"[DEBUG] 公章检测：发现圆形轮廓，面积={area:.0f}，圆形度={circularity:.2f}")
                        return True
        
        # 如果没有找到圆形，但红色占比较高（>1%），也认为可能有公章
        if red_ratio > 0.01:
            print(f"[DEBUG] 公章检测：虽无明显圆形，但红色占比较高 {red_ratio:.2%}，可能有公章")
            return True
        
        print(f"[DEBUG] 公章检测：红色占比 {red_ratio:.2%}，无明显公章特征")
        return False
    
    def _remove_red_aggressive(self, img):
        """激进地去除红色（用于红头区域，因为行间无黑色文字）"""
        # 转换为HSV颜色空间
        hsv = cv2.cvtColor(img, cv2.COLOR_BGR2HSV)
        
        # 使用更宽松的红色HSV范围
        lower_red1 = np.array([0, 20, 20])  # 降低饱和度和亮度阈值
        upper_red1 = np.array([20, 255, 255])  # 扩大色调范围
        lower_red2 = np.array([160, 20, 20])  # 扩大色调范围
        upper_red2 = np.array([180, 255, 255])
        
        # 创建红色掩码
        mask1 = cv2.inRange(hsv, lower_red1, upper_red1)
        mask2 = cv2.inRange(hsv, lower_red2, upper_red2)
        red_mask = mask1 + mask2
        
        # 更强的形态学操作：扩张后闭运算，确保去除干净
        kernel = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (5, 5))
        red_mask = cv2.dilate(red_mask, kernel, iterations=2)  # 扩张
        red_mask = cv2.morphologyEx(red_mask, cv2.MORPH_CLOSE, kernel, iterations=3)
        
        # 将红色区域改为白色
        img[red_mask > 0] = [255, 255, 255]
        
        return img
    
    def _remove_red_protect_black_text(self, img):
        """去除红色但保护黑色文字（用于公章区域）
        策略：精准区分纯红色和黑色文字，只去除纯红色"""
        
        # 转换为浮点数处理
        img_float = img.astype(np.float32)
        b, g, r = cv2.split(img_float)
        
        # 1. 检测红色区域（使用HSV）
        hsv = cv2.cvtColor(img, cv2.COLOR_BGR2HSV)
        lower_red1 = np.array([0, 20, 20])
        upper_red1 = np.array([25, 255, 255])
        lower_red2 = np.array([155, 20, 20])
        upper_red2 = np.array([180, 255, 255])
        
        mask1 = cv2.inRange(hsv, lower_red1, upper_red1)
        mask2 = cv2.inRange(hsv, lower_red2, upper_red2)
        red_mask = (mask1 | mask2).astype(bool)
        
        # 2. 精准区分纯红色和黑色文字
        if np.any(red_mask):
            # 关键区分标准：
            # - 纯红色：R >> (B+G)，且R较高
            # - 黑色文字：RGB都低且接近（中性灰/黑色）
            
            # 计算RGB差异和平均值
            max_val = np.maximum(np.maximum(r, g), b)
            min_val = np.minimum(np.minimum(r, g), b)
            avg_val = (r + g + b) / 3.0
            
            # 判断是否为"接近中性色"（黑色文字的特征）
            # RGB差异小 且 整体较暗
            is_neutral = (max_val - min_val < 60) & (avg_val < 120)
            
            # 判断是否为"纯红色"
            # 条件1：红色通道明显高于蓝绿通道
            # 条件2：红色通道本身较高（>80）
            # 条件3：不是中性色
            is_pure_red = (r > b + 60) & (r > g + 60) & (r > 80) & (~is_neutral)
            
            # 处理红色区域
            red_area = red_mask
            
            # 纯红色：改为白色
            pure_red_pixels = red_area & is_pure_red
            if np.any(pure_red_pixels):
                b[pure_red_pixels] = 255
                g[pure_red_pixels] = 255
                r[pure_red_pixels] = 255
                print(f"[DEBUG] 纯红色→白色: {np.sum(pure_red_pixels)}px")
            
            # 非纯红色（可能有黑色文字）：保留原样或转为中性灰
            has_text = red_area & (~is_pure_red)
            if np.any(has_text):
                # 提取暗度，保留黑色成分
                darkness = np.minimum(np.minimum(b, g), r)
                # 只对很暗的像素（darkness < 100）保留黑色
                very_dark = has_text & (darkness < 100)
                
                if np.any(very_dark):
                    # 转为中性灰/黑色，保留文字
                    dark_val = darkness[very_dark]
                    b[very_dark] = dark_val
                    g[very_dark] = dark_val
                    r[very_dark] = dark_val
                    print(f"[DEBUG] 保留黑色文字: {np.sum(very_dark)}px")
                
                # 不太暗的像素（可能是红色残留）：改为白色
                not_very_dark = has_text & (darkness >= 100)
                if np.any(not_very_dark):
                    b[not_very_dark] = 255
                    g[not_very_dark] = 255
                    r[not_very_dark] = 255
                    print(f"[DEBUG] 浅色红色→白色: {np.sum(not_very_dark)}px")
        
        # 合并通道并转回uint8
        img_result = cv2.merge([b, g, r])
        img_result = np.clip(img_result, 0, 255).astype(np.uint8)
        
        return img_result
    
    def _remove_red_with_header_protection(self, img, protect_ratio=0.63):
        """将图片中红色改为白色，但保护上部指定比例的区域
        只处理有公章的区域，提高效率
        
        Args:
            img: 输入图片
            protect_ratio: 上部保护区域的比例（0.63表示上部63%不处理）
        """
        img_height = img.shape[0]
        protected_height = int(img_height * protect_ratio)
        
        # 只处理保护区域下方的部分（可能有公章）
        if protected_height < img_height:
            bottom_part = img[protected_height:, :].copy()
            # 先检测是否有公章，有才处理
            if self._has_red_seal(bottom_part):
                print("[DEBUG] 检测到公章（下部区域），开始处理...")
                # 使用保护黑色文字的方法处理下部区域
                bottom_part = self._remove_red_protect_black_text(bottom_part)
                img[protected_height:, :] = bottom_part
            else:
                print("[DEBUG] 未检测到公章（下部区域），跳过处理")
        
        return img
    
    def _remove_seal_from_image(self, img, page_num=0):
        """从图片中移除公章（保留下方文字和红头）"""
        # 转换为HSV颜色空间
        hsv = cv2.cvtColor(img, cv2.COLOR_BGR2HSV)
        
        # 扩大红色的HSV范围，提高检测准确性
        lower_red1 = np.array([0, 30, 30])
        upper_red1 = np.array([15, 255, 255])
        lower_red2 = np.array([165, 30, 30])
        upper_red2 = np.array([180, 255, 255])
        
        mask1 = cv2.inRange(hsv, lower_red1, upper_red1)
        mask2 = cv2.inRange(hsv, lower_red2, upper_red2)
        red_mask = mask1 + mask2
        
        # 形态学操作：填补空洞，连接断裂部分
        kernel = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (3, 3))
        red_mask = cv2.morphologyEx(red_mask, cv2.MORPH_CLOSE, kernel, iterations=2)
        red_mask = cv2.morphologyEx(red_mask, cv2.MORPH_OPEN, kernel, iterations=1)
        
        # 红头保护逻辑：只在首页（page_num=0）上部30%区域生效
        red_mask_protected = red_mask
        if page_num == 0:
            # 定义红头保护区域（页面上部30%）
            img_height = img.shape[0]
            header_area_height = int(img_height * 0.3)
            
            # 创建红头保护掩码
            header_protection_mask = np.zeros(img.shape[:2], dtype=np.uint8)
            header_protection_mask[0:header_area_height, :] = 255
            
            # 从红色掩码中排除红头区域
            red_mask_protected = red_mask & (~header_protection_mask)
        
        # 多种方法检测公章区域
        processed = False
        
        # 方法1：检测圆形公章（多种参数组合）
        circles_params = [
            {'param1': 50, 'param2': 30, 'minRadius': 25, 'maxRadius': 250},
            {'param1': 30, 'param2': 20, 'minRadius': 20, 'maxRadius': 300},
            {'param1': 70, 'param2': 40, 'minRadius': 30, 'maxRadius': 200}
        ]
        
        for params in circles_params:
            circles = cv2.HoughCircles(red_mask_protected, cv2.HOUGH_GRADIENT, 1, 50, **params)
            if circles is not None:
                circles = np.round(circles[0, :]).astype("int")
                for (x, y, r) in circles:
                    # 如果是首页，确保圆心不在红头区域
                    if page_num == 0:
                        header_area_height = int(img.shape[0] * 0.3)
                        if y <= header_area_height:
                            continue
                    
                    # 创建圆形掩码，稍微扩大半径以确保完全覆盖
                    mask = np.zeros(img.shape[:2], dtype=np.uint8)
                    cv2.circle(mask, (x, y), int(r * 1.1), 255, -1)
                    
                    # 处理红色像素
                    if page_num == 0:
                        # 首页：不处理红头区域
                        header_protection_mask = np.zeros(img.shape[:2], dtype=np.uint8)
                        header_protection_mask[0:int(img.shape[0] * 0.3), :] = 255
                        red_pixels = (red_mask > 0) & (mask > 0) & (~header_protection_mask > 0)
                    else:
                        # 非首页：可以处理所有红色像素
                        red_pixels = (red_mask > 0) & (mask > 0)
                    
                    img[red_pixels] = [255, 255, 255]
                    processed = True
                break
        
        # 方法2：如果圆形检测失败，使用轮廓检测
        if not processed:
            contours, _ = cv2.findContours(red_mask_protected, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            for contour in contours:
                area = cv2.contourArea(contour)
                # 过滤掉太小的区域（可能是噪点）和太大的区域
                if 500 < area < 50000:
                    # 检查轮廓中心
                    M = cv2.moments(contour)
                    if M["m00"] != 0:
                        cx = int(M["m10"] / M["m00"])
                        cy = int(M["m01"] / M["m00"])
                        
                        # 如果是首页，只处理中心在红头区域外的轮廓
                        if page_num == 0:
                            header_area_height = int(img.shape[0] * 0.3)
                            if cy <= header_area_height:
                                continue
                        
                        # 计算轮廓的圆形度
                        perimeter = cv2.arcLength(contour, True)
                        if perimeter > 0:
                            circularity = 4 * np.pi * area / (perimeter * perimeter)
                            # 如果圆形度较高，认为是公章
                            if circularity > 0.3:
                                # 创建轮廓掩码
                                mask = np.zeros(img.shape[:2], dtype=np.uint8)
                                cv2.fillPoly(mask, [contour], 255)
                                
                                # 稍微扩大掩码区域
                                kernel_expand = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (5, 5))
                                mask = cv2.dilate(mask, kernel_expand, iterations=1)
                                
                                # 处理红色像素
                                if page_num == 0:
                                    # 首页：不处理红头区域
                                    header_protection_mask = np.zeros(img.shape[:2], dtype=np.uint8)
                                    header_protection_mask[0:int(img.shape[0] * 0.3), :] = 255
                                    red_pixels = (red_mask > 0) & (mask > 0) & (~header_protection_mask > 0)
                                else:
                                    # 非首页：可以处理所有红色像素
                                    red_pixels = (red_mask > 0) & (mask > 0)
                                
                                img[red_pixels] = [255, 255, 255]
                                processed = True
        
        # 方法3：如果以上方法都失败，直接处理所有红色区域（但要保护文字和红头）
        if not processed:
            # 检测黑色文字区域
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            _, text_mask = cv2.threshold(gray, 100, 255, cv2.THRESH_BINARY_INV)
            
            # 形态学操作来连接文字
            text_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (3, 1))
            text_mask = cv2.morphologyEx(text_mask, cv2.MORPH_CLOSE, text_kernel, iterations=1)
            
            # 处理红色区域
            if page_num == 0:
                # 首页：不处理与文字重叠且不在红头区域的红色区域
                safe_red_mask = red_mask_protected & (~text_mask)
            else:
                # 非首页：只保护文字，不保护红头区域
                safe_red_mask = red_mask & (~text_mask)
            
            img[safe_red_mask > 0] = [255, 255, 255]
        
        return img
    
    def _convert_text_to_image_pdf(self, input_path, output_path):
        """将文本型PDF转换为图片型PDF"""
        doc = fitz.open(input_path)
        new_doc = fitz.open()

        for page_num in range(len(doc)):
            page = doc[page_num]
            # 保留原始页面的物理尺寸（point），避免用渲染像素数当作页面尺寸导致页面被放大
            orig_width = page.rect.width
            orig_height = page.rect.height

            # 转换为适中分辨率图片，平衡质量和文件大小
            mat = fitz.Matrix(1.5, 1.5)  # 降低分辨率从2.0到1.5
            pix = page.get_pixmap(matrix=mat)
            img_data = pix.tobytes("png")

            # 使用OpenCV进行图像优化
            nparr = np.frombuffer(img_data, np.uint8)
            img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)

            # 转换回PDF页面，使用JPEG压缩
            encode_param = [int(cv2.IMWRITE_JPEG_QUALITY), 85]  # JPEG质量85%
            _, img_encoded = cv2.imencode('.jpg', img, encode_param)
            img_bytes = img_encoded.tobytes()

            # 创建新页面：使用原始物理尺寸，图片会被自动缩放填入，不影响清晰度
            img_rect = fitz.Rect(0, 0, orig_width, orig_height)
            new_page = new_doc.new_page(width=orig_width, height=orig_height)
            new_page.insert_image(img_rect, stream=img_bytes)

        new_doc.save(output_path)
        new_doc.close()
        doc.close()
    
    def _pdf_to_word(self, input_path, output_path):
        """将PDF转换为Word文档，尽可能保持格式"""
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = fitz.open(input_path)
        word_doc = Document()
        
        # 设置页面边距（缩小边距以容纳更多内容）
        sections = word_doc.sections
        for section in sections:
            section.top_margin = Cm(2)      # 上边距2cm
            section.bottom_margin = Cm(2)   # 下边距2cm
            section.left_margin = Cm(2)     # 左边距2cm
            section.right_margin = Cm(2)    # 右边距2cm
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            page_width = page.rect.width
            page_height = page.rect.height
            
            # 使用"dict"模式提取文本，包含格式信息
            text_dict = page.get_text("dict")
            blocks = text_dict.get("blocks", [])
            
            # 按垂直位置排序块
            blocks = sorted([b for b in blocks if b.get("type") == 0], 
                          key=lambda b: b.get("bbox", [0,0,0,0])[1])
            
            for block in blocks:
                # 只处理文本块
                if block.get("type") != 0:
                    continue
                
                lines = block.get("lines", [])
                if not lines:
                    continue
                
                # 获取块的边界框
                bbox = block.get("bbox", [0, 0, page_width, 0])
                block_left = bbox[0]
                block_right = bbox[2]
                block_width = block_right - block_left
                block_center = (block_left + block_right) / 2
                page_center = page_width / 2
                
                # 判断对齐方式（更严格的居中判断）
                alignment = WD_ALIGN_PARAGRAPH.LEFT  # 默认左对齐
                
                # 计算文本块相对于页面中心的偏移
                center_offset = abs(block_center - page_center)
                left_margin = block_left
                right_margin = page_width - block_right
                
                # 只有文本块明显居中且左右边距相近时，才判断为居中
                # 条件：1) 中心偏移小  2) 左右边距差异小  3) 文本宽度不是满页宽
                if (center_offset < 30 and 
                    abs(left_margin - right_margin) < 50 and 
                    block_width < page_width * 0.8):
                    alignment = WD_ALIGN_PARAGRAPH.CENTER
                # 右对齐：右边距很小但左边距很大
                elif right_margin < 50 and left_margin > page_width * 0.3:
                    alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                # 创建段落
                para = word_doc.add_paragraph()
                para.alignment = alignment
                
                # 设置行间距（1.15倍行距，适中）
                para.paragraph_format.line_spacing = 1.15
                # 设置段前段后间距，适当留白
                para.paragraph_format.space_before = Pt(3)
                para.paragraph_format.space_after = Pt(3)
                
                # 为左对齐的段落添加首行缩进（如果PDF中有明显缩进）
                if alignment == WD_ALIGN_PARAGRAPH.LEFT and block_left > 100:
                    # 根据PDF中的实际缩进设置Word缩进
                    indent_ratio = (block_left - 70) / page_width
                    if indent_ratio > 0.05:  # 有明显缩进
                        para.paragraph_format.first_line_indent = Cm(0.75)  # 约2个字符
                
                # 处理每一行
                for line_idx, line in enumerate(lines):
                    spans = line.get("spans", [])
                    
                    # 如果不是第一行，添加换行
                    if line_idx > 0:
                        # 检查是否需要换行（行间距较大说明是新段落）
                        prev_line = lines[line_idx - 1]
                        prev_bbox = prev_line.get("bbox", [0, 0, 0, 0])
                        curr_bbox = line.get("bbox", [0, 0, 0, 0])
                        line_gap = curr_bbox[1] - prev_bbox[3]
                        
                        # 如果行间距很小，在同一段落内；否则可能需要新段落
                        if line_gap > 10:  # 行间距较大，创建新段落
                            para = word_doc.add_paragraph()
                            para.alignment = alignment
                            # 设置行间距（1.15倍行距）
                            para.paragraph_format.line_spacing = 1.15
                            # 设置段前段后间距
                            para.paragraph_format.space_before = Pt(3)
                            para.paragraph_format.space_after = Pt(3)
                            # 新段落也添加首行缩进
                            if alignment == WD_ALIGN_PARAGRAPH.LEFT and block_left > 100:
                                indent_ratio = (block_left - 70) / page_width
                                if indent_ratio > 0.05:
                                    para.paragraph_format.first_line_indent = Cm(0.75)
                    
                    # 处理span
                    for span_idx, span in enumerate(spans):
                        text = span.get("text", "")
                        if not text:
                            continue
                        
                        # 添加文本到段落（不strip，保留原始空格）
                        run = para.add_run(text)
                        
                        # 设置字体
                        font_name = span.get("font", "")
                        if font_name:
                            # 处理中文字体映射
                            font_mapping = {
                                'SimSun': '宋体',
                                'SimHei': '黑体',
                                'FangSong_GB2312': '仿宋',
                                'KaiTi_GB2312': '楷体',
                                'FZXBSJW--GB1-0': '方正小标宋简体',
                            }
                            run.font.name = font_mapping.get(font_name, font_name)
                            # 为中文设置东亚字体
                            try:
                                from docx.oxml.ns import qn
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), 
                                    font_mapping.get(font_name, font_name))
                            except:
                                pass
                        
                        # 设置字体大小
                        size = span.get("size", 0)
                        if size > 0:
                            run.font.size = Pt(size)
                        
                        # 设置颜色
                        color = span.get("color", 0)
                        if isinstance(color, int):
                            if color > 0:
                                r = (color >> 16) & 0xFF
                                g = (color >> 8) & 0xFF
                                b = color & 0xFF
                                run.font.color.rgb = RGBColor(r, g, b)
                        
                        # 设置粗体和斜体
                        flags = span.get("flags", 0)
                        if flags & 2**4:  # 粗体
                            run.font.bold = True
                        if flags & 2**1:  # 斜体
                            run.font.italic = True
            
            # 处理图片块（跳过公章）
            image_blocks = [b for b in text_dict.get("blocks", []) if b.get("type") == 1]
            for block in image_blocks:
                try:
                    bbox = block.get("bbox", None)
                    if not bbox:
                        continue
                    
                    # 从页面提取该区域的图片
                    pix = page.get_pixmap(clip=fitz.Rect(bbox), dpi=150)
                    img_data = pix.tobytes("png")
                    
                    # 检查是否是红色公章，如果是则跳过
                    if self._is_red_seal_image(img_data):
                        print(f"[DEBUG] 跳过公章图片")
                        pix = None
                        continue
                    
                    # 计算图片宽度（保持比例）
                    img_width = bbox[2] - bbox[0]
                    img_height = bbox[3] - bbox[1]
                    
                    # 跳过太小的图片（可能是装饰性元素）
                    if img_width < 30 or img_height < 30:
                        pix = None
                        continue
                    
                    # Word页面宽度约6.5英寸（A4纸去掉边距）
                    doc_width = min(6.5, (img_width / page_width) * 6.5)
                    
                    # 添加图片到Word文档
                    img_stream = io.BytesIO(img_data)
                    word_doc.add_picture(img_stream, width=Inches(doc_width))
                    
                    pix = None
                except Exception as e:
                    print(f"[DEBUG] 处理图片块失败: {e}")
                    continue
            
            # 在每页后添加分页符（除了最后一页）
            if page_num < len(doc) - 1:
                word_doc.add_page_break()
        
        word_doc.save(output_path)
        doc.close()
    
    def _is_red_text(self, color):
        """判断文本颜色是否为红色"""
        # color是一个整数，需要转换为RGB
        if isinstance(color, int):
            r = (color >> 16) & 0xFF
            g = (color >> 8) & 0xFF
            b = color & 0xFF
            
            print(f"[DEBUG] 检测文字颜色 (整数): 原始={color}, RGB=({r}, {g}, {b})")
            
            # 放宽红色判断标准：红色分量明显高于绿色和蓝色
            is_red = r > 100 and r > g * 1.5 and r > b * 1.5
            print(f"[DEBUG] 是否为红色: {is_red}")
            return is_red
        elif isinstance(color, (list, tuple)) and len(color) >= 3:
            # 处理RGB数组格式
            r, g, b = color[0], color[1], color[2]
            # 如果是0-1范围的浮点数，转换为0-255
            if r <= 1.0:
                r_int, g_int, b_int = int(r * 255), int(g * 255), int(b * 255)
                print(f"[DEBUG] 检测文字颜色 (浮点): 原始=({r:.3f}, {g:.3f}, {b:.3f}), RGB=({r_int}, {g_int}, {b_int})")
                is_red = r_int > 100 and r_int > g_int * 1.5 and r_int > b_int * 1.5
            else:
                print(f"[DEBUG] 检测文字颜色 (整数数组): RGB=({int(r)}, {int(g)}, {int(b)})")
                is_red = r > 100 and r > g * 1.5 and r > b * 1.5
            print(f"[DEBUG] 是否为红色: {is_red}")
            return is_red
        print(f"[DEBUG] 无法识别的颜色格式: {color} (类型: {type(color)})")
        return False
    
    def _is_red_seal_image(self, img_data):
        """判断图片是否为红色公章"""
        try:
            nparr = np.frombuffer(img_data, np.uint8)
            img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
            
            if img is None or img.size == 0:
                return False
            
            # 转换为HSV
            hsv = cv2.cvtColor(img, cv2.COLOR_BGR2HSV)
            
            # 检测红色（放宽范围）
            lower_red1 = np.array([0, 30, 50])     # 降低饱和度阈值从50到30
            upper_red1 = np.array([15, 255, 255])  # 扩大色调范围从10到15
            lower_red2 = np.array([165, 30, 50])   # 扩大色调范围从170到165
            upper_red2 = np.array([180, 255, 255])
            
            mask1 = cv2.inRange(hsv, lower_red1, upper_red1)
            mask2 = cv2.inRange(hsv, lower_red2, upper_red2)
            red_mask = mask1 + mask2
            
            # 降低红色像素占比阈值从30%到15%
            total_pixels = img.shape[0] * img.shape[1]
            red_pixels = np.sum(red_mask > 0)
            red_ratio = red_pixels / total_pixels
            
            print(f"[DEBUG] 图片尺寸: {img.shape}, 红色像素占比: {red_ratio:.2%}")
            
            is_seal = red_ratio > 0.15
            print(f"[DEBUG] 是否为红色公章: {is_seal}")
            return is_seal
            
        except Exception as e:
            print(f"[DEBUG] 检测红色公章时出错: {e}")
            return False
    
    def _create_zip(self, base_dir, filenames, zip_path):
        """创建ZIP文件"""
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for filename in filenames:
                file_path = os.path.join(base_dir, filename)
                if os.path.exists(file_path):
                    zipf.write(file_path, filename)